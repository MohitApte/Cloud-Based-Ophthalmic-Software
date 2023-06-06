from tkinter import *
from functools import partial
from tkinter import ttk
import tkinter as tk
import ttkthemes
from reportlab.pdfgen import canvas
from ttkthemes import ThemedStyle
import datetime
import pyautogui as pg
import time
from PIL import Image
import win32print
import win32ui
import win32con
import subprocess
import win32api
import os
from pymongo.mongo_client import MongoClient
from pymongo.server_api import ServerApi
from bson import ObjectId
from docx import Document
from PIL import Image
from tkinter import Tk, Label, Button, filedialog
from pymongo import MongoClient
from PIL import Image
import io
from docx.shared import Inches
from tkinter import messagebox



uri = "mongodb+srv://mohitapte4:<password>@cluster0.xmn1i2w.mongodb.net/?retryWrites=true&w=majority"
client = MongoClient(uri, server_api=ServerApi('1'))
def validateLogin(username, password):
    if username.get() == 'DR' and password.get()=="DR":
        
        main_page()


x = []
medname = []
medtype = []
medadvice = []
days = []
dwm = []
qty = []
img_data = ""

from reportlab.pdfgen import canvas

def draw_multiline_text(canvas, text, x, y, width, height, font_size):
    lines = []
    current_line = ""
    words = text.split()
    max_line_height = 0

    for word in words:
        if canvas.stringWidth(current_line + " " + word, "Helvetica", font_size) < width:
            current_line += " " + word
        else:
            lines.append(current_line)
            current_line = word

    if current_line:
        lines.append(current_line)

    for line in lines:
        line_height = canvas.stringWidth(line, "Helvetica", font_size)
        if line_height > max_line_height:
            max_line_height = line_height

    total_lines = len(lines)
    remaining_height = height - (total_lines * max_line_height)
    y -= max_line_height

    for line in lines:
        canvas.drawString(x, y, line.strip())
        y -= max_line_height

    return remaining_height








def new_patient():
    
    def validateSubmit(mrd,fn,mn,ln,age,sex,address,mob,land,misc):
        today = datetime.date.today()
        
        if ((mob.get().isdigit()) and (land.get().isdigit()) and age.get().isdigit()) and fn.get().isalpha() and mn.get().isalpha() and ln.get().isalpha() and sex.get().isalpha() and address.get().isalnum():
            today_string = today.strftime('%d/%m/%Y')
            document = {"MRD":mrd.get(),"first_name":fn.get(),"middle_name":mn.get(),
                        "last_name":ln.get(),"age":age.get(),"sex":sex.get(),"address":address.get(),"mobile_no":mob.get(),
                        "land_no":land.get(), "misc":misc.get(),
                        'rds':'',
                        'rdc':'',
                        'rda':'',
                        'rdv':'',
                        'rcs':'',
                        'rcc':'',
                        'rca':'',
                        'rcv':'',
                        'rns':'',
                        'rnc':'',
                        'rna':'',
                        'rnv':'',
                        'lds':'',
                        'ldc':'',
                        'lda':'',
                        'ldv':'',
                        'lcs':'',
                        'lcc':'',
                        'lca':'',
                        'lcv':'',
                        'lns':'',
                        'lnc':'',
                        'lna':'',
                        'lnv':'',
                        'ipd':'',
                        'entry1':'',
                        'entry2':'',
                        'entry3':'',
                        'entry4':'',
                        'complaints':'Cheif Complaints:',
                        'examination':'Examination:',
                        'diagnosis':'Diagnosis:',
                        'medicine':'Medicine:',
                        'history':'History:',
                        'advised':'Advised:',
                        'x':[],
                        'medname':[],
                        'medtype':[],
                        'medadvice':[],
                        'days':[],
                        'dwm':[],
                        'qty':[],
                        'img_data':'',
                        'doatxt':'',
                        't1txt':'',
                        'dodtxt':'',
                        't2txt':'',
                        'cftxt':'',
                        'opnotestxt':'',
                        'investigationtxt':'',
                        'postmedicinetxt':'',
                        'surgeryadvisingtxt':'',
                        'adviseondischargetxt':'',
                        'date':''
                        }
            db = client.get_database('patient_data')
            collection = db['patient_name_age']
            collection.insert_one(document)
            
            cursor = collection.find({"$and": [
                                {"MRD": {"$regex": mrd.get(), "$options": "i"}},
                                {"first_name": {"$regex": fn.get(), "$options": "i"}},
                                {"middle_name": {"$regex": mn.get(), "$options": "i"}},
                                {"last_name": {"$regex": ln.get(), "$options": "i"}},
                                {"age": {"$regex": age.get(), "$options": "i"}},
                                {"address": {"$regex": address.get(), "$options": "i"}},
                                {"mobile_no": {"$regex": mob.get(), "$options": "i"}},
                                {"land_no": {"$regex": land.get(), "$options": "i"}},
                                {"misc": {"$regex": misc.get(), "$options": "i"}}
                            ]})
            data = [doc for doc in cursor]
            for doc in data:
                values = [str(v) for v in doc.values()]
            
                
            tab1.destroy()
            patient_selected(values)
        else:
            messagebox.showerror("Error", "An error occurred!")
            return
        
    tab1 = Toplevel(root)
    mrdLabel = ttk.Label(tab1, text="MRD Number").grid(row=0, column=0)
    mrd = StringVar()
    mrdEntry = ttk.Entry(tab1, textvariable=mrd).grid(row=0, column=1)
    import random
    today = datetime.date.today()
    new_day = today.day
    mrd_str = str(new_day)+str(random.randint(100000, 999999))
    mrd.set(mrd_str)
    
    
    fnLabel = ttk.Label(tab1, text="First Name").grid(row=2, column=0)
    fn = StringVar()
    fnEntry = ttk.Entry(tab1, textvariable=fn).grid(row=2, column=1) 

    mnLabel = ttk.Label(tab1, text="Middle Name").grid(row=4, column=0)
    mn = StringVar()
    mnEntry = ttk.Entry(tab1, textvariable=mn).grid(row=4, column=1) 

    lnLabel = ttk.Label(tab1, text="Last Name").grid(row=6, column=0)
    ln = StringVar()
    lnEntry = ttk.Entry(tab1, textvariable=ln).grid(row=6, column=1) 

    ageLabel = ttk.Label(tab1, text="Age").grid(row=8, column=0)
    age = StringVar()
    ageEntry = ttk.Entry(tab1, textvariable=age).grid(row=8, column=1) 

    sexLabel = ttk.Label(tab1, text="Sex").grid(row=10, column=0)
    sex = StringVar()
    ttk.Radiobutton(tab1,variable=sex, text="Male",value="Male", command=None).grid(row=10, column=1)
    ttk.Radiobutton(tab1,variable=sex, text="Female",value="Female", command=None).grid(row=10, column=2)

    addressLabel = ttk.Label(tab1, text="Address").grid(row=12, column=0)
    address = StringVar()
    addressEntry = ttk.Entry(tab1, textvariable=address).grid(row=12, column=1) 

    mobLabel = ttk.Label(tab1, text="Mobile Number").grid(row=14, column=0)
    mob = StringVar()
    mobEntry = ttk.Entry(tab1, textvariable=mob).grid(row=14, column=1) 

    landLabel = ttk.Label(tab1, text="Landline Number").grid(row=16, column=0)
    land = StringVar()
    landEntry = ttk.Entry(tab1, textvariable=land).grid(row=16, column=1)


    miscLabel = ttk.Label(tab1, text="Miscellaneous").grid(row=18, column=0)
    misc = StringVar()
    miscEntry = ttk.Entry(tab1, textvariable=misc).grid(row=18, column=1)

    validateSubmit = partial(validateSubmit, mrd,fn,mn,ln,age,sex,address,mob,land,misc)
    submitButton = ttk.Button(tab1, text="Submit", command=validateSubmit).grid(row=20, column=0)
    

def old_patient():
    
    def validateSubmit(mrd,fn,mn,ln,age,sex,address,mob,land,misc):

            
            def on_tree_select(event):
                item = event.widget.selection()[0]
                values = event.widget.item(item)['values']
                tab2.destroy()
                patient_selected(values)
                
            db = client.get_database('patient_data')
            collection = db['patient_name_age']
            
            cursor = collection.find({"$and": [
                                {"MRD": {"$regex": mrd.get(), "$options": "i"}},   
                                {"first_name": {"$regex": fn.get(), "$options": "i"}},
                                {"middle_name": {"$regex": mn.get(), "$options": "i"}},
                                {"last_name": {"$regex": ln.get(), "$options": "i"}},
                                {"age": {"$regex": age.get(), "$options": "i"}},
                                {"sex": {"$regex": sex.get()}},
                                {"address": {"$regex": address.get(), "$options": "i"}},
                                {"mobile_no": {"$regex": mob.get(), "$options": "i"}},
                                {"land_no": {"$regex": land.get(), "$options": "i"}},
                                {"misc": {"$regex": misc.get(), "$options": "i"}}
                            ]})
            data = [doc for doc in cursor]
            tree_frame = ttk.Frame(tab2)
            tree_frame.grid(column=3, row=22, sticky='nsew')
    
            # Create a tkinter window and Treeview widget to display the JSON data
            tree = ttk.Treeview(tab2, columns=list(data[0].keys()), show='headings')
            for key in data[0].keys():
                tree.heading(key, text=key)
            mrd_data={}
            for doc in data:
                values = [str(v) for v in doc.values()]
                if values[1] not in mrd_data:    
                    mrd_data[values[1]] = 1
                    tree.insert('', 'end', values=values)
                    
                    
            tree_scroll = ttk.Scrollbar(tree_frame, orient='vertical', command=tree.yview)
            tree.configure(yscrollcommand=tree_scroll.set)
            
            tree.grid(column=3, row=22, sticky='nsew')
            tree_scroll.grid(column=4, row=22, sticky='ns')
            # Bind the treeview widget to a function that will be called when a row is selected
            tree.bind('<<TreeviewSelect>>', on_tree_select)
       
            

        

    
    tab2 = Toplevel(root)
    screen_width = tab2.winfo_screenwidth()
    screen_height = tab2.winfo_screenheight()
    tab2.geometry("%dx%d" % (screen_width, screen_height))
    tab2.title("Old Patient")
    mrdLabel = ttk.Label(tab2, text="MRD Number").grid(row=0, column=0)
    mrd = StringVar()
    mrdEntry = ttk.Entry(tab2, textvariable=mrd).grid(row=0, column=1)

    fnLabel = ttk.Label(tab2, text="First Name").grid(row=2, column=0)
    fn = StringVar()
    fnEntry = ttk.Entry(tab2, textvariable=fn).grid(row=2, column=1) 

    mnLabel = ttk.Label(tab2, text="Middle Name").grid(row=4, column=0)
    mn = StringVar()
    mnEntry = ttk.Entry(tab2, textvariable=mn).grid(row=4, column=1) 

    lnLabel = ttk.Label(tab2, text="Last Name").grid(row=6, column=0)
    ln = StringVar()
    lnEntry = ttk.Entry(tab2, textvariable=ln).grid(row=6, column=1) 

    ageLabel = ttk.Label(tab2, text="Age").grid(row=8, column=0)
    age = StringVar()
    ageEntry = ttk.Entry(tab2, textvariable=age).grid(row=8, column=1) 
    
    sexLabel = ttk.Label(tab2, text="Sex").grid(row=10, column=0)
    sex = StringVar()
    ttk.Radiobutton(tab2,variable=sex, text="Male",value="Male", command=None).grid(row=10, column=1)
    ttk.Radiobutton(tab2,variable=sex, text="Female",value="Female", command=None).grid(row=10, column=2)

    addressLabel = ttk.Label(tab2, text="Address").grid(row=12, column=0)
    address = StringVar()
    addressEntry = ttk.Entry(tab2, textvariable=address).grid(row=12, column=1) 

    mobLabel = ttk.Label(tab2, text="Mobile Number").grid(row=14, column=0)
    mob = StringVar()
    mobEntry = ttk.Entry(tab2, textvariable=mob).grid(row=14, column=1) 

    landLabel = ttk.Label(tab2, text="Landline Number").grid(row=16, column=0)
    land = StringVar()
    landEntry = ttk.Entry(tab2, textvariable=land).grid(row=16, column=1)


    miscLabel = ttk.Label(tab2, text="Miscellaneous").grid(row=18, column=0)
    misc = StringVar()
    miscEntry = ttk.Entry(tab2, textvariable=misc).grid(row=18, column=1)

    validateSubmit = partial(validateSubmit, mrd,fn,mn,ln,age,sex, address,mob,land,misc)
    submitButton = ttk.Button(tab2, text="Search", command=validateSubmit).grid(row=20, column=0)

   

def main_page():
    app = Toplevel(root)
    screen_width = app.winfo_screenwidth()
    screen_height = app.winfo_screenheight()
    app.geometry("%dx%d" % (screen_width, screen_height))
    tabControl = ttk.Notebook(app)
    tab1 = ttk.Frame(tabControl)
    tab2 = ttk.Frame(tabControl)
    tabControl.add(tab1, text ='Out Patient Department')
    tabControl.add(tab2, text ='In Patient Department')
    tabControl.pack(expand = 1, fill ="both")
    ttk.Button(tab1, text="New Patient", command=new_patient).grid(row=1, column=0) 
    
    ttk.Button(tab1, text="Old Patient", command=old_patient).grid(row=1, column=1)
    
    global patient_selected
    global cur_pat
    def patient_selected(doc):
        global cur_pat
        cur_pat = doc
        db = client.get_database('patient_data')
        collection = db['patient_name_age']
        
        cursor = collection.find({"$and": [
                            {"first_name": {"$regex": str(doc[2]), "$options": "i"}},
                            {"middle_name": {"$regex": str(doc[3]), "$options": "i"}},
                            {"last_name": {"$regex": str(doc[4]), "$options": "i"}},
                        ]})
        document = cursor[0]
        name = doc[2] + ' ' + doc[3] + " " + doc[4]
        today = datetime.date.today()
        today_string = today.strftime('%d/%m/%Y')
        patient_info_frame = ttk.LabelFrame(tab1, text = "Patient Information")
        patient_info_frame.grid(row = 2, column = 0)
        
        ttk.Label(patient_info_frame, text="Name: "+ name, borderwidth=3, relief="ridge").grid(row = 3, column= 0)
        ttk.Label(patient_info_frame, text="Age: " +str(doc[5]), borderwidth=3, relief="ridge").grid(row = 3, column= 1)
        ttk.Label(patient_info_frame, text="Sex: "+str(doc[6]), borderwidth=3, relief="ridge").grid(row = 3, column= 2)
        ttk.Label(patient_info_frame, text="Mob: "+str(doc[8]), borderwidth=3, relief="ridge").grid(row = 3, column= 3)
        ttk.Label(patient_info_frame, text="Date: "+today_string, borderwidth=3, relief="ridge").grid(row = 3, column= 4)
        
        patient_detail_frame = ttk.LabelFrame(tab1, text = "Patient Details")
        patient_detail_frame.grid(row = 4, column = 0)
        def chief_complaints(event):
            complaints = Toplevel(root)
            complaints.geometry("500x500")
            
        
        def history(event):
            history = Toplevel(root)
            history.geometry("500x500")
        
        def exam(event):
            exam = Toplevel(root)
            exam.geometry("500x500")
            
        def diagram(event):
            diagram = Toplevel(root)
            diagram.geometry("500x500")
            global img_data
            def select_image():
                global img_data
                filename = filedialog.askopenfilename(initialdir="/", title="Select Image", filetypes=(("Image files", "*.jpg *.jpeg *.png"), ("All files", "*.*")))
                if filename:
                    image = Image.open(filename)
                    image.show()
            
                    # Store the image in MongoDB
                    with open(filename, "rb") as f:
                        img_data = f.read()
                        #print(img_data)
                    
                    
             
                    
            select_image()
           

                
                
            
            
            
            
        def diagnosis(event):
            diagnosis = Toplevel(root)
            diagnosis.geometry("500x500")
            
        def advised(event):
            advised = Toplevel(root)
            advised.geometry("500x500")
            
        
        
            
        
            
        def chief_medicine(event):
            medicine = Toplevel(root)
            
            medicine.attributes("-fullscreen", True)
            screen_width = medicine.winfo_screenwidth()
            screen_height = medicine.winfo_screenheight()
            medicine.geometry("%dx%d" % (screen_width, screen_height))
            
            patient_info_frame = ttk.LabelFrame(medicine, text = "Patient Information")
            patient_info_frame.grid(row = 0, column = 0)
            
            medicine_frame = ttk.LabelFrame(medicine, text = "Medicine")
            medicine_frame.grid(row = 2, column = 0)
            ttk.Label(patient_info_frame, text="Name: "+ name, borderwidth=3, relief="ridge").grid(row = 9, column= 1)
            ttk.Label(patient_info_frame, text="Age: " +str(doc[5]), borderwidth=3, relief="ridge").grid(row = 9, column= 2)
            ttk.Label(patient_info_frame, text="Sex: "+str(doc[6]), borderwidth=3, relief="ridge").grid(row = 9, column= 3)
            ttk.Label(patient_info_frame, text="Mob: "+str(doc[8]), borderwidth=3, relief="ridge").grid(row = 9, column= 4)
            ttk.Label(patient_info_frame, text="Date: "+today_string+"             ", borderwidth=3, relief="ridge").grid(row = 9, column= 5)
            
            
            ttk.Label(medicine_frame, text="Root Of Administration").grid(row = 1, column= 1)

            
            ttk.Label(medicine_frame, text="Medicine Name").grid(row = 1, column= 5)
            ttk.Label(medicine_frame, text="Type").grid(row = 1, column= 6)
            ttk.Label(medicine_frame, text="Advice").grid(row = 1, column= 7)
            ttk.Label(medicine_frame, text="Days").grid(row = 1, column= 8)
            ttk.Label(medicine_frame, text="DWM").grid(row = 1, column= 9)
            ttk.Label(medicine_frame, text="Qty").grid(row = 1, column= 10)
            
            
            
            global currow
            currow = 1 
            vari = {}
            global x
            global medname
            global medtype
            global medadvice
            global days
            global dwm
            global qty
            
            
            
            def add_more():
                global currow
                currow+=1
                vari[str(currow)+"xtxt"] = ttk.Combobox(medicine_frame, values=["Right Eye", "Left Eye", "Both Eyes", "Oral", "IM", "IV"]) 
                vari[str(currow)+"xtxt"].grid(row=currow, column=1)
                
                

                
                vari[str(currow)+"mednametxt"] = Entry(medicine_frame, width=20)
                vari[str(currow)+"mednametxt"].grid(row=currow, column=5)
                
                
                vari[str(currow)+"typetxt"] = Entry(medicine_frame, width=4)
                vari[str(currow)+"typetxt"].grid(row=currow, column=6)
                
                vari[str(currow)+"medadvicetxt"] = Entry(medicine_frame, width=10)
                vari[str(currow)+"medadvicetxt"].grid(row=currow, column=7)
                
                
                vari[str(currow)+"daystxt"] = Entry(medicine_frame, width=4)
                vari[str(currow)+"daystxt"].grid(row=currow, column=8)
                
                
                vari[str(currow)+"dwmtxt"] = Entry(medicine_frame, width=2)
                vari[str(currow)+"dwmtxt"].grid(row=currow, column=9)
                
                
                vari[str(currow)+"qtytxt"] = Entry(medicine_frame, width=2)
                vari[str(currow)+"qtytxt"].grid(row=currow, column=10)
                
                
                try:    
                    vari[str(currow)+"xtxt"].insert(END, x[currow-2])
                    vari[str(currow)+"mednametxt"].insert(END, medname[currow-2])
                    vari[str(currow)+"typetxt"].insert(END, medtype[currow-2])
                    vari[str(currow)+"medadvicetxt"].insert(END, medadvice[currow-2])
                    vari[str(currow)+"daystxt"].insert(END, days[currow-2])
                    vari[str(currow)+"dwmtxt"].insert(END, dwm[currow-2])
                    vari[str(currow)+"qtytxt"].insert(END, qty[currow-2])
                except:
                    pass
     
        
                return
            
            
            for i in range(len(x)):
                add_more()
            
            
            
            def save():
                x.clear()

                medname.clear()
                medtype.clear()
                medadvice.clear()
                days.clear()
                dwm.clear()
                qty.clear()
                for i in range(currow-1):
                    x.append(vari[str(i+2)+"xtxt"].get())
                    medname.append(vari[str(i+2)+"mednametxt"].get())
                    medtype.append(vari[str(i+2)+"typetxt"].get())
                    medadvice.append(vari[str(i+2)+"medadvicetxt"].get())
                    days.append(vari[str(i+2)+"daystxt"].get())
                    dwm.append(vari[str(i+2)+"dwmtxt"].get())
                    qty.append(vari[str(i+2)+"qtytxt"].get())
                    
                print(x)   
                print(medname)
            
            
            button = ttk.Button(medicine, text="New", command=add_more)
            button.grid(row = 12, column= 2, sticky=tk.S)
            
            button = ttk.Button(medicine, text="Save", command=save)
            button.grid(row = 12, column= 3, sticky=tk.S)
                
            
            def screenshot():
                pdf = canvas.Canvas("medicine.pdf")
                pdf.drawString(100, 800, name)
                pdf.drawString(250, 800, str(doc[5]))
                pdf.drawString(280, 800, str(doc[6]))
                pdf.drawString(330, 800, today_string)
                
                row = 750
                for i in range(currow - 1):
                
                    pdf.drawString(100, row, medname[i])
                    pdf.drawString(170, row, medtype[i])
                    pdf.drawString(270, row, medadvice[i])
                    pdf.drawString(370, row, "in "+x[i])
                    pdf.drawString(470, row, days[i]+" days")
                    pdf.drawString(520, row, "("+qty[i]+")")
                    
                    
                    pdf.drawString(100, row-30, "----------------------------------------------------------------------------------")
                    row -= 100
                pdf.save()
                
                filename = "medicine.pdf"
                
                if os.name == "posix":  # for macOS or Linux
                    os.system("open " + filename)
                elif os.name == "nt":  # for Windows
                    os.system("start " + filename)
            
            def exit_window():
                    medicine.destroy() 
            
            
            exit_button = ttk.Button(medicine, text="Exit", command=exit_window)
            exit_button.grid(row = 12, column= 5, sticky=tk.S)   
                
                
            button = ttk.Button(medicine_frame, text="Print", command=screenshot)
            button.grid(row = 12, column= 4, sticky=tk.S)
            
            
            
            
        def prescription(event):
            prescription = Toplevel(root)
            prescription.attributes("-fullscreen", True)
            screen_width = prescription.winfo_screenwidth()
            screen_height = prescription.winfo_screenheight()
            prescription.geometry("%dx%d" % (screen_width, screen_height))

            
            def insertValue(value, text_field, num_win):
                text_field.insert(END, value)
                num_win.destroy()
            
            def nums(event, field_txt):
                num_win = Toplevel(root)
                num_win.geometry("700x700")
                button1 = ttk.Button(num_win, text="0.75", command=lambda val=str(0.75): insertValue(val,field_txt,num_win))
                button1.grid(row=0, column=0)
                button2 = ttk.Button(num_win, text="0.1", command=lambda val=str(0.1): insertValue(val,field_txt,num_win))
                button2.grid(row=0, column=1)
                button3 = ttk.Button(num_win, text="0.2", command=lambda val=str(0.2): insertValue(val,field_txt,num_win))
                button3.grid(row=0, column=2)
                button4 = ttk.Button(num_win, text="0.25", command=lambda val=str(0.25): insertValue(val,field_txt,num_win))
                button4.grid(row=0, column=3)
                button5 = ttk.Button(num_win, text="0.3", command=lambda val=str(0.3): insertValue(val,field_txt,num_win))
                button5.grid(row=0, column=4)
                button6 = ttk.Button(num_win, text="0.35", command=lambda val=str(0.35): insertValue(val,field_txt,num_win))
                button6.grid(row=0, column=5)
                button7 = ttk.Button(num_win, text="0.4", command=lambda val=str(0.4): insertValue(val,field_txt,num_win))
                button7.grid(row=1, column=0)
                button8 = ttk.Button(num_win, text="0.45", command=lambda val=str(0.45): insertValue(val,field_txt,num_win))
                button8.grid(row=1, column=1)
                button9 = ttk.Button(num_win, text="0.5", command=lambda val=str(0.5): insertValue(val,field_txt,num_win))
                button9.grid(row=1, column=2)
                
            patient_info_frame = ttk.LabelFrame(prescription, text = "Patient Information")
            patient_info_frame.grid(row = 0, column = 0)
            
            glass_prescription_frame = ttk.LabelFrame(prescription, text = "Glass Prescription")
            glass_prescription_frame.grid(row = 2, column = 0)
            
            
            def exit():
                prescription.destroy()
            
            
            
            exit_button = ttk.Button(prescription, text="Exit", command=exit)
            exit_button.grid(row = 12, column= 5, sticky=tk.S)
            
            rds = document['rds']
            rdc = document['rdc']
            rda = document['rda']
            rdv = document['rdv']
            rcs = document['rcs']
            rcc = document['rcc']
            rca = document['rca']
            rcv = document['rcv']
            rns = document['rns']
            rnc = document['rnc']
            rna = document['rna']
            rnv = document['rnv']
            lds = document['lds']
            ldc = document['ldc']
            lda = document['lda']
            ldv = document['ldv']
            lcs = document['lcs']
            lcc = document['lcc']
            lca = document['lca']
            lcv = document['lcv']
            lns = document['lns']
            lnc = document['lnc']
            lna = document['lna']
            lnv = document['lnv']
            ipd = document['ipd']
            entry1txt = document['entry1']
            entry2txt = document['entry2']
            entry3txt = document['entry3']
            entry4txt = document['entry4']
              
            
            ttk.Label(patient_info_frame, text="Name: "+ name, borderwidth=3, relief="ridge").grid(row = 9, column= 1)
            ttk.Label(patient_info_frame, text="Age: " +str(doc[5]), borderwidth=3, relief="ridge").grid(row = 9, column= 2)
            ttk.Label(patient_info_frame, text="Sex: "+str(doc[6]), borderwidth=3, relief="ridge").grid(row = 9, column= 3)
            ttk.Label(patient_info_frame, text="Mob: "+str(doc[8]), borderwidth=3, relief="ridge").grid(row = 9, column= 4)
            ttk.Label(patient_info_frame, text="Date: "+today_string+"             ", borderwidth=3, relief="ridge").grid(row = 9, column= 5)
            
            
            
            
            
            ttk.Label(glass_prescription_frame, text="Right Eye").grid(row = 0, column= 1,columnspan=4)    
            ttk.Label(glass_prescription_frame, text="Left Eye").grid(row = 0, column= 5, columnspan=4)
            
            ttk.Label(glass_prescription_frame, text="Dist.").grid(row = 2, column= 0)
            ttk.Label(glass_prescription_frame, text="Computer").grid(row = 3, column= 0)
            ttk.Label(glass_prescription_frame, text="Near").grid(row = 4, column= 0)
            
            
            ttk.Label(glass_prescription_frame, text="Spl.").grid(row = 1, column= 1)
            ttk.Label(glass_prescription_frame, text="Cyl.").grid(row = 1, column= 2)
            ttk.Label(glass_prescription_frame, text="Axis").grid(row = 1, column= 3)
            ttk.Label(glass_prescription_frame, text="Vision").grid(row = 1, column= 4)
            
            ttk.Label(glass_prescription_frame, text="Spl.").grid(row = 1, column= 5)
            ttk.Label(glass_prescription_frame, text="Cyl.").grid(row = 1, column= 6)
            ttk.Label(glass_prescription_frame, text="Axis").grid(row = 1, column= 7)
            ttk.Label(glass_prescription_frame, text="Vision").grid(row = 1, column= 8)
            
            
            
            rdstxt = Entry(glass_prescription_frame, width=5)
            rdstxt.grid(row=2, column=1)
            rdstxt.bind("<Double-Button-1>", lambda event: nums(event, rdstxt), add="+")

            
            rdctxt = Entry(glass_prescription_frame, width=5)
            rdctxt.grid(row=2, column=2)
            rdctxt.bind("<Double-Button-1>", lambda event: nums(event, rdctxt), add="+")
            
            rdatxt = Entry(glass_prescription_frame, width=5)
            rdatxt.grid(row=2, column=3)
            rdatxt.bind("<Double-Button-1>", lambda event: nums(event, rdatxt), add="+")
            
            rdvtxt = Entry(glass_prescription_frame, width=5)
            rdvtxt.grid(row=2, column=4)
            rdvtxt.bind("<Double-Button-1>", lambda event: nums(event, rdvtxt), add="+")
            
            rcstxt = Entry(glass_prescription_frame, width=5)
            rcstxt.grid(row=3, column=1)
            rcstxt.bind("<Double-Button-1>", lambda event: nums(event, rcstxt), add="+")
            
            
            rcctxt = Entry(glass_prescription_frame, width=5)
            rcctxt.grid(row=3, column=2)
            rcctxt.bind("<Double-Button-1>", lambda event: nums(event, rcctxt), add="+")
            
            rcatxt = Entry(glass_prescription_frame, width=5)
            rcatxt.grid(row=3, column=3)
            rcatxt.bind("<Double-Button-1>", lambda event: nums(event, rcatxt), add="+")
            
            rcvtxt = Entry(glass_prescription_frame, width=5)
            rcvtxt.grid(row=3, column=4)
            rcvtxt.bind("<Double-Button-1>", lambda event: nums(event, rcvtxt), add="+")
            
            
            rnstxt = Entry(glass_prescription_frame, width=5)
            rnstxt.grid(row=4, column=1)
            rnstxt.bind("<Double-Button-1>", lambda event: nums(event, rnstxt), add="+")
            
            rnctxt = Entry(glass_prescription_frame, width=5)
            rnctxt.grid(row=4, column=2)
            rnctxt.bind("<Double-Button-1>", lambda event: nums(event, rnctxt), add="+")
            
            rnatxt = Entry(glass_prescription_frame, width=5)
            rnatxt.grid(row=4, column=3)
            rnatxt.bind("<Double-Button-1>", lambda event: nums(event, rnatxt), add="+")
            
            rnvtxt = Entry(glass_prescription_frame, width=5)
            rnvtxt.grid(row=4, column=4)
            rnvtxt.bind("<Double-Button-1>", lambda event: nums(event, rnvtxt), add="+")
            
            
            
            ldstxt = Entry(glass_prescription_frame, width=5)
            ldstxt.grid(row=2, column=5)
            ldstxt.bind("<Double-Button-1>", lambda event: nums(event, ldstxt), add="+")
            
            ldctxt = Entry(glass_prescription_frame, width=5)
            ldctxt.grid(row=2, column=6)
            ldctxt.bind("<Double-Button-1>", lambda event: nums(event, ldctxt), add="+")
            
            ldatxt = Entry(glass_prescription_frame, width=5)
            ldatxt.grid(row=2, column=7)
            ldatxt.bind("<Double-Button-1>", lambda event: nums(event, ldatxt), add="+")
            
            ldvtxt = Entry(glass_prescription_frame, width=5)
            ldvtxt.grid(row=2, column=8)
            ldvtxt.bind("<Double-Button-1>", lambda event: nums(event, ldvtxt), add="+")
            
            
            lcstxt = Entry(glass_prescription_frame, width=5)
            lcstxt.grid(row=3, column=5)
            lcstxt.bind("<Double-Button-1>", lambda event: nums(event, lcstxt), add="+")
            
            lcctxt = Entry(glass_prescription_frame, width=5)
            lcctxt.grid(row=3, column=6)
            lcctxt.bind("<Double-Button-1>", lambda event: nums(event, lcctxt), add="+")
            
            lcatxt = Entry(glass_prescription_frame, width=5)
            lcatxt.grid(row=3, column=7)
            lcatxt.bind("<Double-Button-1>", lambda event: nums(event, lcatxt), add="+")
            
            lcvtxt = Entry(glass_prescription_frame, width=5)
            lcvtxt.grid(row=3, column=8)
            lcvtxt.bind("<Double-Button-1>", lambda event: nums(event, lcvtxt), add="+")
            
            
            lnstxt = Entry(glass_prescription_frame, width=5)
            lnstxt.grid(row=4, column=5)
            lnstxt.bind("<Double-Button-1>", lambda event: nums(event, lnstxt), add="+")
            
            lnctxt = Entry(glass_prescription_frame, width=5)
            lnctxt.grid(row=4, column=6)
            lnctxt.bind("<Double-Button-1>", lambda event: nums(event, lnctxt), add="+")
            
            lnatxt = Entry(glass_prescription_frame, width=5)
            lnatxt.grid(row=4, column=7)
            lnatxt.bind("<Double-Button-1>", lambda event: nums(event, lnatxt), add="+")
            
            lnvtxt = Entry(glass_prescription_frame, width=5)
            lnvtxt.grid(row=4, column=8)
            lnvtxt.bind("<Double-Button-1>", lambda event: nums(event, lnvtxt), add="+")
            
            
            glass_details_frame = ttk.LabelFrame(prescription, text = "Glass Details")
            glass_details_frame.grid(row = 5, column = 0)
            
            ttk.Label(glass_details_frame, text="IPD:").grid(row = 5, column= 0)
            ipdtxt = Entry(glass_details_frame)
            ipdtxt.grid(row=5, column=1)
            ipdtxt.bind("<Double-Button-1>", lambda event: nums(event, ipdtxt), add="+")
            
            
            ttk.Label(glass_details_frame, text="Purpose: ").grid(row = 7, column= 0)
            global entry1
            entry1 = ttk.Combobox(glass_details_frame, values=["Constant Use", "For Near Vision Only", "For Distant Vision Only"])
            entry1.grid(row = 7, column= 1)
            entry2 = ttk.Combobox(glass_details_frame, values=["White Glass", "Photo Grey", "High Index Glass", "Anti-Glare Coating", "Goggles"])
            entry2.grid(row = 7, column= 2)
            
            
            ttk.Label(glass_details_frame, text="Remark: ").grid(row = 8, column= 0)
            entry3 = ttk.Combobox(glass_details_frame, values=["Kryptok Bifocal Glass", "Executive Bifocal Glass", "Progressive Glass"])
            entry3.grid(row = 8, column= 1)
            entry4 = ttk.Combobox(glass_details_frame, values=["Change Right Glass Only","Change Left Glass Only", "Change Both Glasses"])
            entry4.grid(row = 8, column= 2)
            
            
            rdstxt.insert(END, rds)
            rdctxt.insert(END, rdc)
            rdatxt.insert(END, rda)
            rdvtxt.insert(END, rdv)
            rcstxt.insert(END, rcs)
            rcctxt.insert(END, rcc)
            rcatxt.insert(END, rca)
            rcvtxt.insert(END, rcv)
            rnstxt.insert(END, rns)
            rnctxt.insert(END, rnc)
            rnatxt.insert(END, rna)
            rnvtxt.insert(END, rnv)
            ldstxt.insert(END, lds)
            ldctxt.insert(END, ldc)
            ldatxt.insert(END, lda)
            ldvtxt.insert(END, ldv)
            lcstxt.insert(END, lcs)
            lcctxt.insert(END, lcc)
            lcatxt.insert(END, lca)
            lcvtxt.insert(END, lcv)
            lnstxt.insert(END, lns)
            lnctxt.insert(END, lnc)
            lnatxt.insert(END, lna)
            lnvtxt.insert(END, lnv)
            ipdtxt.insert(END, ipd)
            entry1.insert(END, entry1txt)
            entry2.insert(END, entry2txt)
            entry3.insert(END, entry3txt)
            entry4.insert(END, entry4txt)
                        
            
            
            def save():
                filter = {"$and": [
                                    {"first_name": {"$regex": str(doc[2]), "$options": "i"}},
                                    {"middle_name": {"$regex": str(doc[3]), "$options": "i"}},
                                    {"last_name": {"$regex": str(doc[4]), "$options": "i"}},
                                ]}


                new_values = {'$set': {'rds': rdstxt.get(),
                       'rdc': rdctxt.get(),
                       'rda': rdatxt.get(),
                       'rdv': rdvtxt.get(),
                       'rcs': rcstxt.get(),
                       'rcc': rcctxt.get(),
                       'rca': rcatxt.get(),
                       'rcv': rcvtxt.get(),
                       'rns': rnstxt.get(),
                       'rnc': rnctxt.get(),
                       'rna': rnatxt.get(),
                       'rnv': rnvtxt.get(),
                       'lds': ldstxt.get(),
                       'ldc': ldctxt.get(),
                       'lda': ldatxt.get(),
                       'ldv': ldvtxt.get(),
                       'lcs': lcstxt.get(),
                       'lcc': lcctxt.get(),
                       'lca': lcatxt.get(),
                       'lcv': lcvtxt.get(),
                       'lns': lnstxt.get(),
                       'lnc': lnctxt.get(),
                       'lna': lnatxt.get(),
                       'lnv': lnvtxt.get(),
                       'ipd': ipdtxt.get(),
                       'entry1': entry1.get(),
                       'entry2': entry2.get(),
                       'entry3': entry3.get(),
                       'entry4': entry4.get()}}


                result = collection.update_one(filter, new_values)
            
            
            button = ttk.Button(prescription, text="              Save", command=save)
            button.grid(row = 12, column= 4, sticky=tk.S)
            
            
            
            
            
            def screenshot():
                random = int(time.time())
                filename = "C:/Users/mohit/Desktop/screenshot.png"
                ss = pg.screenshot(filename)
                
                
                # Crop the image using the calculated dimensions
                cropped_image = ss.crop((0, 0, 650, 500))
                cropped_image.save('cropped_screenshot.png')
                
                opened_image = Image.open('cropped_screenshot.png')
                file_path = "cropped_screenshot.png"
                win32api.ShellExecute(
                    0,
                    "print",
                    file_path,
                    f'/d:"{win32print.GetDefaultPrinter()}"',
                    os.path.dirname(file_path),
                    0
                )
            

                
                
                
            button = ttk.Button(prescription, text="              Print", command=screenshot)
            button.grid(row = 12, column= 3, sticky=tk.S)
            
        complaints = document['complaints']
        examination = document['examination']
        diagnosis = document['diagnosis']
        medicine = document['medicine']
        history = document['history']
        advised = document['advised'] 
            
        ttk.Label(patient_detail_frame, text="COMPLAINTS").grid(row=3, column=0)
        complaintxt = Text(patient_detail_frame, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        complaintxt.grid(row=4, column=0)
        
        complaintxt.insert(END, "complaints")
        complaintxt.bind("<Double-Button-1>", chief_complaints, add="+")
        
        
        
        
        
        
        
        ttk.Label(patient_detail_frame, text="HISTORY").grid(row=5, column=0)
        historytxt = Text(patient_detail_frame, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        historytxt.grid(row=6, column=0)
        
        historytxt.insert(END, "history")
        historytxt.bind("<Double-Button-1>", history, add="+")
        
        
        
        ttk.Label(patient_detail_frame, text="EXAMINATION").grid(row=3, column=1)
        examtxt = Text(patient_detail_frame, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        examtxt.grid(row=4, column=1)
        
        examtxt.insert(END, "examination")
        examtxt.bind("<Double-Button-1>", exam, add="+")
        
        
        ttk.Label(patient_detail_frame, text="DIAGRAM").grid(row=5, column=1)
        diagramtxt = Text(patient_detail_frame, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        diagramtxt.grid(row=6, column=1)
        
        diagramtxt.insert(END, "Diagram")
        diagramtxt.bind("<Double-Button-1>", diagram, add="+")
            
        
        ttk.Label(patient_detail_frame, text="DIAGNOSIS").grid(row=3, column=2)
        diagnosistxt = Text(patient_detail_frame, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        diagnosistxt.grid(row=4, column=2)
        
        diagnosistxt.insert(END, "diagnosis")
        diagnosistxt.bind("<Double-Button-1>", diagnosis, add="+")
        
        
        ttk.Label(patient_detail_frame, text="ADVISED").grid(row=5, column=2)
        advisedtxt = Text(patient_detail_frame, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        advisedtxt.grid(row=6, column=2)
        
        advisedtxt.insert(END, "advised")
        advisedtxt.bind("<Double-Button-1>", advised, add="+")
        
        
        ttk.Label(patient_detail_frame, text="MEDICINE").grid(row=3, column=3)
        medicinetxt = Text(patient_detail_frame, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        medicinetxt.grid(row=4, column=3)
        
        medicinetxt.insert(END, "medicine")
        medicinetxt.bind("<Double-Button-1>", chief_medicine, add="+")
        
        
        ttk.Label(patient_detail_frame, text="PRESCRIPTION").grid(row=5, column=3)
        prescriptiontxt = Text(patient_detail_frame, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        prescriptiontxt.grid(row=6, column=3)
        
        prescriptiontxt.insert(END, "Prescription:")
        prescriptiontxt.bind("<Double-Button-1>", prescription, add="+")
        
        def new_save():
            filter = {"$and": [
                                {"first_name": {"$regex": str(doc[2]), "$options": "i"}},
                                {"middle_name": {"$regex": str(doc[3]), "$options": "i"}},
                                {"last_name": {"$regex": str(doc[4]), "$options": "i"}},
                            ]}


            global img_data
            result = collection.find_one(filter)
            source_dict = dict(result)
            new_id = ObjectId()

            source_dict['complaints'] = complaintxt.get("1.0",'end-1c')
            source_dict['examination'] = examtxt.get("1.0",'end-1c')
            source_dict['diagnosis'] = diagnosistxt.get("1.0",'end-1c')
            source_dict['medicine'] = medicinetxt.get("1.0",'end-1c')
            source_dict['history'] = historytxt.get("1.0",'end-1c')
            source_dict['advised'] = advisedtxt.get("1.0",'end-1c')
            source_dict['date'] = today_string
            source_dict['_id'] = new_id
            
            
            
            source_dict['doatxt'] = doatxt.get()
            source_dict['t1txt'] = t1txt.get()
            source_dict['dodtxt'] = dodtxt.get()
            source_dict['t2txt'] = t2txt.get()
            source_dict['cftxt'] = cftxt.get("1.0",'end-1c')
            source_dict['opnotestxt'] = opnotestxt.get("1.0",'end-1c')
            source_dict['investigationtxt'] = investigationtxt.get("1.0",'end-1c')
            source_dict['postmedicinetxt'] = postmedicinetxt.get("1.0",'end-1c')
            source_dict['surgeryadvisingtxt'] = surgeryadvisingtxt.get("1.0",'end-1c')
            source_dict['adviseondischargetxt'] = adviseondischargetxt.get("1.0",'end-1c')
            
            
            
            
            
            
            source_dict['x'] = x
            source_dict['medname'] = medname
            source_dict['medtype'] = medtype
            source_dict['medadvice'] = medadvice
            source_dict['days'] = days
            source_dict['dwm'] = dwm
            source_dict['qty'] = qty
            source_dict['img_data'] = img_data
            print(img_data)            
            collection.insert_one(source_dict)
        
        
        new_button = ttk.Button(patient_detail_frame, text="Save", command=new_save)
        new_button.grid(row = 12, column= 4, sticky=tk.S)
        
        
        def view_history():
            view_hist = Toplevel(root)
            view_hist.geometry("1000x1000")
            global cur_pat
            
            
            def retrive_data(date):
                db = client.get_database('patient_data')
                collection = db['patient_name_age']
                new_col = collection.find({"$and": [
                                    {"first_name": {"$regex": str(cur_pat[2]), "$options": "i"}},
                                    {"middle_name": {"$regex": str(cur_pat[3]), "$options": "i"}},
                                    {"last_name": {"$regex": str(cur_pat[4]), "$options": "i"}},
                                    {"date": {"$regex": str(date), "$options": "i"}}
                                ]})
                
                
                
                for document in new_col:

                    string = ""
                    string += str(document['complaints'])
                    string+="\n\n"
                    string += str(document['history'])
                    string+="\n\n"
                    string += str(document['examination'])
                    string+="\n\n"
                    string += str(document['diagnosis'])
                    string+="\n\n"
                    string += str(document['advised'])
                    string+="\n\n"
                    
                    for i in range(len(document['x'])):
                        
                        string += document['medname'][i] + "            "
                        string += document['medtype'][i] + "            "
                        string += document['medadvice'][i] + "            "
                        string += "in " + document['x'][i] + "         "
                        string += document['days'][i]+" days" + "            "
                        string += "("+document['qty'][i]+")"
                        string += "\n"
                        string += "--------------------------------"
                        string += "\n"
                    
                    
                    output_path = "output.docx"

                    doc = Document()
                    doc.add_paragraph(string)
                    
                        

                    bin_data  = document['img_data']
                    print(bin_data)
                    image_stream = io.BytesIO(bin_data)
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()
                    
                    # Insert the image into the run
                    run.add_picture(image_stream, width=Inches(4))  
                    
                    doc.save(output_path)

                    filename = "output.docx"
                    
                    if os.name == "posix":  # for macOS or Linux
                        os.system("open " + filename)
                    elif os.name == "nt":  # for Windows
                        os.system("start " + filename)
                

                
            db = client.get_database('patient_data')
            collection = db['patient_name_age']
            
            cursor = collection.find({"$and": [
                                {"first_name": {"$regex": str(cur_pat[2]), "$options": "i"}},
                                {"middle_name": {"$regex": str(cur_pat[3]), "$options": "i"}},
                                {"last_name": {"$regex": str(cur_pat[4]), "$options": "i"}},
                            ]})
            data = [doc for doc in cursor]
            new_dict = {}
            row = 0
            
            for doc in data:
                values = [str(v) for v in doc.values()]
                ttk.Button(view_hist, text=str(values[-1]), command=lambda date=values[-1]: retrive_data(date)).grid(row = row, column= 1, sticky=tk.S)
                row += 1
                
                    
          
            
            
        
        hist_button = ttk.Button(patient_detail_frame, text="Complete History", command=view_history)
        hist_button.grid(row = 13, column= 4, sticky=tk.S)
        
        
        
        #Tab 2 IPD
        
        patient_info_ipd = ttk.LabelFrame(tab2, text = "Patient Information")
        patient_info_ipd.grid(row = 1, column = 0)
        patient_discharge_ipd = ttk.LabelFrame(tab2, text = "Discharge Summary")
        patient_discharge_ipd.grid(row = 2, column = 0)
        
        ttk.Label(patient_discharge_ipd, text="Date of Admission", borderwidth=3, relief="ridge").grid(row = 1, column= 0)
        ttk.Label(patient_discharge_ipd, text="Time", borderwidth=3, relief="ridge").grid(row = 1, column= 4)
        ttk.Label(patient_discharge_ipd, text="Date of Discharge", borderwidth=3, relief="ridge").grid(row = 2, column= 0)
        ttk.Label(patient_discharge_ipd, text="Time", borderwidth=3, relief="ridge").grid(row = 2, column= 4)
        
        doatxt = Entry(patient_discharge_ipd, width=10)
        doatxt.grid(row=1, column=1)
        
        t1txt = Entry(patient_discharge_ipd, width=10)
        t1txt.grid(row=1, column=5)
        
        
        dodtxt = Entry(patient_discharge_ipd, width=10)
        dodtxt.grid(row=2, column=1)
        
        t2txt = Entry(patient_discharge_ipd, width=10)
        t2txt.grid(row=2, column=5)
        
        
        ttk.Label(patient_info_ipd, text="Name: "+ name, borderwidth=3, relief="ridge").grid(row = 3, column= 0)
        ttk.Label(patient_info_ipd, text="Age: " +str(doc[5]), borderwidth=3, relief="ridge").grid(row = 3, column= 1)
        ttk.Label(patient_info_ipd, text="Sex: "+str(doc[6]), borderwidth=3, relief="ridge").grid(row = 3, column= 2)
        ttk.Label(patient_info_ipd, text="Mob: "+str(doc[8]), borderwidth=3, relief="ridge").grid(row = 3, column= 3)
        ttk.Label(patient_info_ipd, text="Date: "+today_string, borderwidth=3, relief="ridge").grid(row = 3, column= 4)
        
        patient_detail_ipd = ttk.LabelFrame(tab2, text = "Patient Details IPD")
        patient_detail_ipd.grid(row = 4, column = 0)
        
        
        ttk.Label(patient_detail_ipd, text="Clinical Findings").grid(row=3, column=0)
        cftxt = Text(patient_detail_ipd, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        cftxt.grid(row=4, column=0)
        
        cftxt.insert(END, "Clinical Findings")
        
        
        ttk.Label(patient_detail_ipd, text="Operation Notes").grid(row=5, column=0)
        opnotestxt = Text(patient_detail_ipd, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        opnotestxt.grid(row=6, column=0)
        
        opnotestxt.insert(END, "Operation Notes")
        
        
        
        ttk.Label(patient_detail_ipd, text="Invesitgation").grid(row=3, column=1)
        investigationtxt = Text(patient_detail_ipd, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        investigationtxt.grid(row=4, column=1)
        
        investigationtxt.insert(END, "Invesitgation")
        
        
        ttk.Label(patient_detail_ipd, text="Post Operative Medicines").grid(row=5, column=1)
        postmedicinetxt = Text(patient_detail_ipd, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        postmedicinetxt.grid(row=6, column=1)
        
        postmedicinetxt.insert(END, "Post Operative Medicines")

        

        ttk.Label(patient_detail_ipd, text="Surgery Advising").grid(row=3, column=2)
        surgeryadvisingtxt = Text(patient_detail_ipd, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        surgeryadvisingtxt.grid(row=4, column=2)
        
        surgeryadvisingtxt.insert(END, "Surgery Advising")
        
        ttk.Label(patient_detail_ipd, text="Advice on Discharge").grid(row=5, column=2)
        adviseondischargetxt = Text(patient_detail_ipd, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        adviseondischargetxt.grid(row=6, column=2)
        
        adviseondischargetxt.insert(END, "Advice on Discharge")
        
        ipd_button = ttk.Button(patient_detail_ipd, text="Save", command=new_save)
        ipd_button.grid(row = 12, column= 4, sticky=tk.S)
        
        
        def view_ipd_history():
            view_hist = Toplevel(root)
            view_hist.geometry("1000x1000")
            global cur_pat
            
            
            def retrive_data(date):
                db = client.get_database('patient_data')
                collection = db['patient_name_age']
                new_col = collection.find({"$and": [
                                    {"first_name": {"$regex": str(cur_pat[2]), "$options": "i"}},
                                    {"middle_name": {"$regex": str(cur_pat[3]), "$options": "i"}},
                                    {"last_name": {"$regex": str(cur_pat[4]), "$options": "i"}},
                                    {"date": {"$regex": str(date), "$options": "i"}}
                                ]})
                
                
                
                for document in new_col:

                    string = ""
                    
                    string += str(document['doatxt']) + "  "

                    
                    string += str(document['t1txt'])
                    string+="\n\n"
                    
                    string += str(document['dodtxt']) + "  "

                    
                    string += str(document['t2txt'])
                    string+="\n\n"
                    
                    
                    string += str(document['cftxt'])
                    string+="\n\n"
                    string += str(document['opnotestxt'])
                    string+="\n\n"
                    string += str(document['investigationtxt'])
                    string+="\n\n"
                    string += str(document['postmedicinetxt'])
                    string+="\n\n"
                    string += str(document['surgeryadvisingtxt'])
                    string+="\n\n"
                    string += str(document['adviseondischargetxt'])
                    string+="\n\n"
                    
                    output_path = "ipd.docx"

                    doc = Document()
                    doc.add_paragraph(string)
                    
                    
                    doc.save(output_path)

                    filename = "ipd.docx"
                    
                    if os.name == "posix":  # for macOS or Linux
                        os.system("open " + filename)
                    elif os.name == "nt":  # for Windows
                        os.system("start " + filename)
                

                
            db = client.get_database('patient_data')
            collection = db['patient_name_age']
            
            cursor = collection.find({"$and": [
                                {"first_name": {"$regex": str(cur_pat[2]), "$options": "i"}},
                                {"middle_name": {"$regex": str(cur_pat[3]), "$options": "i"}},
                                {"last_name": {"$regex": str(cur_pat[4]), "$options": "i"}},
                            ]})
            data = [doc for doc in cursor]
            new_dict = {}
            row = 0
            
            for doc in data:
                values = [str(v) for v in doc.values()]
                ttk.Button(view_hist, text=str(values[-1]), command=lambda date=values[-1]: retrive_data(date)).grid(row = row, column= 1, sticky=tk.S)
                row += 1
                
                    
          
            
            
        
        ipd_hist_button = ttk.Button(patient_detail_ipd, text="IPD History", command=view_ipd_history)
        ipd_hist_button.grid(row = 13, column= 4, sticky=tk.S)
        
        
        
        
        
        
        
        
        
        
        
        
        def print_every():
            pdf = canvas.Canvas("hello.pdf")
            pdf.drawString(100, 800, name)
            pdf.drawString(100, 700, complaintxt.get("1.0",'end-1c'))
            pdf.drawString(100, 600, diagnosistxt.get("1.0",'end-1c'))
            pdf.drawString(100, 500, examtxt.get("1.0",'end-1c'))
            pdf.drawString(100, 400, medicinetxt.get("1.0",'end-1c'))
            pdf.drawString(100, 300, advisedtxt.get("1.0",'end-1c'))
            pdf.save()
            
            filename = "hello.pdf"
            
            if os.name == "posix":  # for macOS or Linux
                os.system("open " + filename)
            elif os.name == "nt":  # for Windows
                os.system("start " + filename)
        
        
        
        
        
root = Tk()
root.geometry("400x150")
root.title("Ophthalmic Software")


# Set the initial theme
root.tk.call("source", "azure.tcl")
root.tk.call("set_theme", "light")


usernameLabel = ttk.Label(root, text="User Name").grid(row=0, column=0)
username = StringVar()
usernameEntry = ttk.Entry(root, textvariable=username).grid(row=0, column=1)  

#password label and password entry box
passwordLabel = ttk.Label(root,text="Password").grid(row=1, column=0)  
password = StringVar()
passwordEntry = ttk.Entry(root, textvariable=password, show='*').grid(row=1, column=1)  

validateLogin = partial(validateLogin, username, password)
#login button
loginButton = ttk.Button(root, text="Login", command=validateLogin).grid(row=4, column=0)  


root.mainloop()





